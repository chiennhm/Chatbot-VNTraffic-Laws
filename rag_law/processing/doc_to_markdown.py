# -*- coding: utf-8 -*-
"""
Convert .doc files to Markdown preserving legal document structure
Structures: Phần, Chương, Mục, Điều, Khoản, Điểm
"""

import argparse
import os
import re
import sys
from pathlib import Path

try:
    from dotenv import load_dotenv
    load_dotenv()
except ImportError:
    pass

# Configuration from environment
RAG_LAW_DIR = Path(__file__).parent.parent
PROJECT_ROOT = RAG_LAW_DIR.parent
DOC_DIR = os.getenv("DOC_DIR", str(PROJECT_ROOT / "documents_doc"))
MARKDOWN_DIR = os.getenv("MARKDOWN_DIR", str(PROJECT_ROOT / "documents_markdown"))

def check_dependencies():
    """Check required packages"""
    try:
        import docx
        return True
    except ImportError:
        # Try python-docx
        try:
            from docx import Document
            return True
        except ImportError:
            pass
    
    print("Missing dependency. Install with:")
    print("  pip install python-docx")
    print("\nFor .doc files (old format), also need:")
    print("  pip install pywin32  (Windows only)")
    return False


def convert_doc_to_docx(doc_path: str, output_dir: str) -> str:
    """Convert old .doc to .docx using Word COM (Windows only)"""
    try:
        import win32com.client
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        doc = word.Documents.Open(doc_path)
        docx_path = os.path.join(output_dir, os.path.basename(doc_path).replace('.doc', '.docx'))
        doc.SaveAs2(docx_path, FileFormat=16)  # 16 = docx format
        doc.Close()
        word.Quit()
        
        return docx_path
    except Exception as e:
        print(f"  Warning: Cannot convert .doc: {e}")
        return None


def extract_text_from_docx(docx_path: str) -> str:
    """Extract text from .docx file"""
    from docx import Document
    
    doc = Document(docx_path)
    paragraphs = []
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)
    
    # Also extract from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                paragraphs.append(row_text)
    
    return '\n'.join(paragraphs)


def format_to_markdown(text: str, doc_name: str) -> str:
    """
    Convert plain text to Markdown with legal document structure.
    Recognizes: Phần, Chương, Mục, Điều, Khoản, Điểm
    """
    lines = text.split('\n')
    md_lines = []
    
    # Document title (first non-empty line or from filename)
    title = doc_name.replace('.doc', '').replace('.docx', '').replace('_', ' ')
    md_lines.append(f"# {title}\n")
    
    for line in lines:
        line = line.strip()
        if not line:
            md_lines.append("")
            continue
        
        # PHẦN (Part) - # 
        if re.match(r'^PHẦN\s+[IVXLCDM\d]+', line, re.IGNORECASE):
            md_lines.append(f"\n# {line}\n")
            continue
        
        # CHƯƠNG (Chapter) - ##
        if re.match(r'^CHƯƠNG\s+[IVXLCDM\d]+', line, re.IGNORECASE):
            md_lines.append(f"\n## {line}\n")
            continue
        
        # MỤC (Section) - ###
        if re.match(r'^MỤC\s+\d+', line, re.IGNORECASE):
            md_lines.append(f"\n### {line}\n")
            continue
        
        # Điều (Article) - ####
        if re.match(r'^Điều\s+\d+\.?', line):
            md_lines.append(f"\n#### {line}\n")
            continue
        
        # Khoản (Clause) - numbered list
        khoan_match = re.match(r'^(\d+)\.\s+(.+)', line)
        if khoan_match:
            num, content = khoan_match.groups()
            md_lines.append(f"\n**{num}.** {content}")
            continue
        
        # Điểm (Point) - lettered list
        diem_match = re.match(r'^([a-zđ])\)\s+(.+)', line, re.IGNORECASE)
        if diem_match:
            letter, content = diem_match.groups()
            md_lines.append(f"   - **{letter})** {content}")
            continue
        
        # Sub-points with dash
        if line.startswith('-'):
            md_lines.append(f"      {line}")
            continue
        
        # Regular paragraph
        md_lines.append(line)
    
    return '\n'.join(md_lines)


def clean_text(text: str) -> str:
    """Clean up extracted text"""
    # Remove multiple spaces
    text = re.sub(r' +', ' ', text)
    
    # Remove multiple newlines
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Fix common OCR/extraction issues
    text = text.replace('\r\n', '\n')
    text = text.replace('\r', '\n')
    
    return text.strip()


def convert_doc_to_markdown(doc_path: str, output_dir: str, temp_dir: str) -> bool:
    """Convert a single .doc file to Markdown"""
    filename = os.path.basename(doc_path)
    md_filename = filename.replace('.doc', '.md').replace('.docx', '.md')
    md_path = os.path.join(output_dir, md_filename)
    
    try:
        # If .doc (not .docx), convert first
        if doc_path.lower().endswith('.doc') and not doc_path.lower().endswith('.docx'):
            docx_path = convert_doc_to_docx(doc_path, temp_dir)
            if not docx_path:
                return False
        else:
            docx_path = doc_path
        
        # Extract text
        text = extract_text_from_docx(docx_path)
        
        # Clean text
        text = clean_text(text)
        
        # Convert to Markdown
        markdown = format_to_markdown(text, filename)
        
        # Save
        with open(md_path, 'w', encoding='utf-8') as f:
            f.write(markdown)
        
        # Clean up temp file
        if docx_path != doc_path and os.path.exists(docx_path):
            os.remove(docx_path)
        
        return True
        
    except Exception as e:
        print(f"Error: {e}")
        return False


def main() -> int:
    parser = argparse.ArgumentParser(description="Convert .doc files to Markdown")
    parser.add_argument("--input", "-i", default=DOC_DIR, help="Input directory with .doc files")
    parser.add_argument("--output", "-o", default=MARKDOWN_DIR, help="Output directory for .md files")
    args = parser.parse_args()
    
    input_dir = args.input
    output_dir = args.output
    temp_dir = os.path.join(os.path.dirname(output_dir), 'temp_docx')
    
    # Create directories
    os.makedirs(output_dir, exist_ok=True)
    os.makedirs(temp_dir, exist_ok=True)
    
    print("=" * 60)
    print("DOC TO MARKDOWN CONVERTER")
    print("=" * 60)
    print(f"Input:  {input_dir}")
    print(f"Output: {output_dir}")
    
    if not os.path.exists(input_dir):
        print(f"Error: Input directory not found: {input_dir}")
        return 1
    
    # Find .doc files
    doc_files = [f for f in os.listdir(input_dir) if f.lower().endswith('.doc')]
    print(f"\nFound {len(doc_files)} .doc files")
    
    if not doc_files:
        print("No .doc files found")
        return 0
    
    # Convert each file
    success = 0
    failed = 0
    
    for i, doc_file in enumerate(doc_files, 1):
        doc_path = os.path.join(input_dir, doc_file)
        print(f"\n[{i}/{len(doc_files)}] {doc_file}")
        
        if convert_doc_to_markdown(doc_path, output_dir, temp_dir):
            print(f"Converted")
            success += 1
        else:
            failed += 1
    
    # Cleanup temp dir
    try:
        os.rmdir(temp_dir)
    except:
        pass
    
    print("\n" + "=" * 60)
    print("CONVERSION COMPLETE")
    print("=" * 60)
    print(f"Success: {success}")
    print(f"Failed:  {failed}")
    print(f"Output:  {output_dir}")
    return 0


if __name__ == "__main__":
    if not check_dependencies():
        sys.exit(1)
    sys.exit(main())

